from openai import OpenAI
import streamlit as st
from dotenv import load_dotenv
import os
import PyPDF2
import docx
import io

# Load API key from the .env file
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

def extract_text_from_pdf(pdf_file):
    """Extract text from the uploaded PDF file."""
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page_num in range(len(pdf_reader.pages)):
        text += pdf_reader.pages[page_num].extract_text()
    return text

def extract_text_from_txt(txt_file):
    """Extract text from the uploaded TXT file."""
    return txt_file.getvalue().decode("utf-8")

def extract_text_from_docx(docx_file):
    """Extract text from the uploaded DOCX file."""
    doc = docx.Document(docx_file)
    full_text = []
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    
    return '\n'.join(full_text)

def extract_text_from_doc(doc_file):
    """
    Note: This is a fallback for old .doc files. 
    It's recommended to use .docx format for better compatibility.
    """
    try:
        # Try to read it as a .docx file first
        return extract_text_from_docx(doc_file)
    except:
        st.warning("Old .doc format detected. For better results, please convert to .docx format.")
        return "Error: Unable to process .doc format. Please convert to .docx and try again."

def extract_text(file, file_type):
    """Extract text based on file type."""
    try:
        if file_type == "pdf":
            return extract_text_from_pdf(file)
        elif file_type == "txt":
            return extract_text_from_txt(file)
        elif file_type == "docx":
            return extract_text_from_docx(file)
        elif file_type == "doc":
            return extract_text_from_doc(file)
        return ""
    except Exception as e:
        st.error(f"Error extracting text: {str(e)}")
        return ""

def ask_openai(prompt, max_tokens=300):
    """Send a prompt to OpenAI using the latest Chat API."""
    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",  # or "gpt-4" if you have access
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=max_tokens
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        st.error(f"Error: {e}")
        return "Something went wrong. Please try again."

def process_large_text(text, max_chunk_size=2000):
    """Split text into chunks if it's too large."""
    chunks = []
    words = text.split()
    current_chunk = []
    current_length = 0
    
    for word in words:
        current_length += len(word) + 1  # +1 for space
        if current_length > max_chunk_size:
            chunks.append(' '.join(current_chunk))
            current_chunk = [word]
            current_length = len(word)
        else:
            current_chunk.append(word)
    
    if current_chunk:
        chunks.append(' '.join(current_chunk))
    
    return chunks

# Streamlit UI
st.title("Document Genie: Ask Questions About Your Document")

# File upload
uploaded_file = st.file_uploader("Choose a document file!", type=["pdf", "txt", "docx", "doc"])

if uploaded_file is not None:
    # Get file type
    file_type = uploaded_file.name.split('.')[-1].lower()
    
    with st.spinner(f"Extracting text from the {file_type.upper()} file..."):
        document_text = extract_text(uploaded_file, file_type)

    if document_text:
        st.success(f"{file_type.upper()} content extracted successfully!")
        
        # Add a word count
        word_count = len(document_text.split())
        st.info(f"Document contains approximately {word_count} words")
        
        # Show extracted text in an expandable section
        with st.expander("View Document Content"):
            st.text_area("Extracted Content", document_text, height=200)

        # Question input and processing
        question = st.text_input("Ask a question about the document:")
        
        if st.button("Find answer"):
            with st.spinner("Finding the answer..."):
                # Process large documents in chunks if necessary
                if len(document_text) > 4000:  # OpenAI's token limit consideration
                    chunks = process_large_text(document_text)
                    answers = []
                    
                    progress_bar = st.progress(0)
                    for i, chunk in enumerate(chunks):
                        prompt = f"Based on this part of the document, answer the question. If the answer cannot be found in this part, just respond with 'NOT_FOUND'.\n\nText: {chunk}\n\nQuestion: {question}"
                        chunk_answer = ask_openai(prompt, max_tokens=150)
                        if chunk_answer != "NOT_FOUND":
                            answers.append(chunk_answer)
                        progress_bar.progress((i + 1) / len(chunks))
                    
                    if answers:
                        final_prompt = f"Combine and summarize these answers to the question '{question}':\n\n{' '.join(answers)}"
                        final_answer = ask_openai(final_prompt, max_tokens=300)
                    else:
                        final_answer = "I couldn't find a relevant answer in the document."
                else:
                    prompt = f"Based on the following text, answer the question. If the answer cannot be found in the text, say so.\n\nText: {document_text}\n\nQuestion: {question}"
                    final_answer = ask_openai(prompt)
                
                # Display the answer in a more prominent way
                st.markdown("### Answer:")
                st.markdown(f"> {final_answer}")
                
                # Add a button to ask a follow-up question
                if st.button("Ask a follow-up question"):
                    st.session_state.previous_answer = final_answer
                    st.text_input("What else would you like to know?")
    else:
        st.warning(f"The {file_type.upper()} file seems empty or couldn't be processed. Please upload a valid file.")
else:
    st.info("Upload a PDF, TXT, DOCX, or DOC file to start.")

# Add some helpful tips in the sidebar
with st.sidebar:
    st.markdown("### Supported File Types:")
    st.markdown("""
    - PDF (.pdf)
    - Text files (.txt)
    - Word documents (.docx, .doc*)
    
    *Note: For .doc files, converting to .docx format is recommended for better results.
    """)
    
    st.markdown("### Tips for better results:")
    st.markdown("""
    1. Make sure your document is clearly formatted
    2. For PDFs, ensure they are searchable (not scanned images)
    3. Ask specific questions rather than general ones
    4. Break down complex questions into simpler ones
    5. For large documents, be patient as processing may take longer
    """)
    
    st.markdown("### Document Size Limits:")
    st.markdown("""
    - The application will automatically handle large documents by splitting them into manageable chunks
    - For best results, try to keep documents under 10,000 words
    """)